import re
import copy
from docx import Document
from docx.oxml.ns import qn
from typing import List, Optional
from io import BytesIO

from rag.nlp import rag_tokenizer


# ── Temizleme ────────────────────────────────────────────────────────────────

def is_strikethrough(run) -> bool:
    try:
        return bool(run.font.strike) or bool(run.font.double_strike)
    except:
        return False


def clean_variants(text: str) -> str:
    if not text:
        return text
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:[^)]*\)\s*', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:.*$', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\(?\s*Yürürlük\s*:.*$', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\b\d+\s*md\.\s*Yürürlük\s*:.*$', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_clean_text(paragraph) -> Optional[str]:
    all_parts = []
    for run in paragraph.runs:
        text = run.text
        if not text.strip():
            continue
        all_parts.append({
            'text': text,
            'is_strike': is_strikethrough(run)
        })
    clean_parts = [p['text'] for p in all_parts if not p['is_strike']]
    if not clean_parts:
        return None
    combined = ''.join(clean_parts)
    cleaned = clean_variants(combined)
    return cleaned.strip() if cleaned.strip() else None


# ── Madde numarası tespiti ────────────────────────────────────────────────────

MADDE_NO_PATTERN = re.compile(r'^(\d+(?:\.\d+)*)\.\s*')


def get_word_numbering(paragraph, doc) -> Optional[str]:
    """Word otomatik numaralama key'ini döndür."""
    try:
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is None:
            return None
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return None
        ilvl_el = numPr.find(qn('w:ilvl'))
        numId_el = numPr.find(qn('w:numId'))
        if ilvl_el is None or numId_el is None:
            return None
        ilvl = ilvl_el.get(qn('w:val'), '0')
        numId = numId_el.get(qn('w:val'), '0')
        return f"numId:{numId},ilvl:{ilvl}"
    except Exception:
        return None


def detect_madde_no_from_text(text: str) -> Optional[str]:
    m = MADDE_NO_PATTERN.match(text.strip())
    if m:
        return m.group(1)
    return None


# ── Ana chunker sınıfı ────────────────────────────────────────────────────────

LISTE_ADI = "EK-4/F AYAKTA TEDAVİDE SAĞLIK RAPORU (Uzman Hekim Raporu/Sağlık Kurulu Raporu) İLE VERİLEBİLECEK İLAÇLAR LİSTESİ"


class Ek4fDocx:
    """
    EK-4/F listesi için madde bazlı chunker.
    Her madde ayrı bir chunk olarak üretilir.
    """

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        doc = Document(filename) if not binary else Document(BytesIO(binary))

        chunks = []
        mevcut_madde_no = None
        mevcut_satirlar = []
        content_started = False
        num_counters = {}

        def kaydet_madde():
            if mevcut_madde_no is None or not mevcut_satirlar:
                return
            icerik = ' '.join(mevcut_satirlar)
            icerik = re.sub(r'\s+', ' ', icerik).strip()
            if icerik:
                chunk = f"Liste Adı: {LISTE_ADI}, Madde {mevcut_madde_no}: {icerik}"
                chunks.append(chunk)

        for para in doc.paragraphs:
            raw_text = para.text.strip()
            if not raw_text:
                continue

            line = extract_clean_text(para)
            numbering_key = get_word_numbering(para, doc)

            madde_no = None
            icerik_kismi = None

            if numbering_key:
                if numbering_key not in num_counters:
                    num_counters[numbering_key] = 0
                num_counters[numbering_key] += 1
                madde_no = str(num_counters[numbering_key])
                icerik_kismi = line.strip() if line else ""
            elif line:
                madde_no_text = detect_madde_no_from_text(line)
                if madde_no_text:
                    madde_no = madde_no_text
                    icerik_kismi = MADDE_NO_PATTERN.sub('', line.strip()).strip()

            if not content_started:
                if madde_no:
                    content_started = True
                else:
                    continue

            if madde_no:
                kaydet_madde()
                mevcut_madde_no = madde_no
                mevcut_satirlar = [icerik_kismi] if icerik_kismi else []
            else:
                if mevcut_madde_no is not None and line:
                    mevcut_satirlar.append(line.strip())

        kaydet_madde()
        return chunks


# ── RAG entegrasyon fonksiyonu ────────────────────────────────────────────────

def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Turkish", callback=None, **kwargs):
    """
    RAG Entegrasyon Fonksiyonu

    EK-4/F listesi için madde bazlı chunking.
    Her madde ayrı bir chunk olarak üretilir.
    """

    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "EK-4F isleniyor (madde bazli chunking)...")
        rawChunks = Ek4fDocx()(filename, binary, from_page, to_page)
        callback(0.7, "EK-4F chunking tamamlandi.")

        result = []
        for chunkContent in rawChunks:
            d = copy.deepcopy(doc)
            d["content_with_weight"] = chunkContent
            d["content_ltks"] = rag_tokenizer.tokenize(chunkContent)
            d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
            result.append(d)

        callback(0.9, f"{len(result)} chunk olusturuldu.")
        return result

    else:
        raise NotImplementedError("EK-4F chunker sadece .docx dosyalarini destekler")


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], callback=dummy)