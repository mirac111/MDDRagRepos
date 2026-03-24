import re
import copy
from docx import Document
from typing import List, Dict, Optional
from io import BytesIO

from common.constants import ParserType
from rag.nlp import rag_tokenizer


class RegexPatterns:
    """Filtreleme için regex pattern'ları."""
    
    # İçindekiler satırları
    TOC = re.compile(r"\.{4,}\s+\d+\s*$")
    
    # Bölüm başlıkları (içeriğin başlangıcını tespit için)
    BOLUM = re.compile(
        r"^(BİRİNCİ|İKİNCİ|ÜÇÜNCÜ|DÖRDÜNCÜ|BEŞİNCİ|ALTINCI|YEDİNCİ|SEKİZİNCİ|DOKUZUNCU|ONUNCU)\s+BÖLÜM$"
    )
    
    # Sayfa numaraları
    PAGE_NUMBER = re.compile(r'^\s*\d+\s*$')
    
    # SUT Ekleri (örnek: "EK-2/A-2", "EK-1/A", "EK-3/B-1")
    SUT_EK = re.compile(
        r'EK-\s*[1-4]\s*/\s*[A-ZÇĞİÖŞÜ](?:\s*-\s*\d+)?', 
        re.IGNORECASE
    )


class DocxStyleAnalyzer:
    """DOCX formatlamasını analiz eder."""
    
    @staticmethod
    def is_strikethrough(run) -> bool:
        """FIX: None değil bool döndür"""
        try:
            strike = run.font.strike
            double_strike = run.font.double_strike
            return bool(strike) or bool(double_strike)
        except:
            return False
    
    @staticmethod
    def is_red(run) -> bool:
        """FIX: Her zaman bool döndür"""
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                r, g, b = rgb[0] if len(rgb) > 0 else 0, rgb[1] if len(rgb) > 1 else 0, rgb[2] if len(rgb) > 2 else 0
                return bool(r > 150 and g < 100 and b < 100)
            return False
        except:
            return False


class TextExtractor:
    """DOCX'ten temiz metin çıkarır."""
    
    @staticmethod
    def clean_variants(text: str) -> str:
        """
        (Mülga:...), (Değişik:...), (Ek:...), (Yürürlük:...) gibi 
        varyant bilgilerini metinden temizler.
        """
        if not text:
            return text
        
        # Boşlukları normalize et
        text = re.sub(r'\s+', ' ', text)
        
        # Kapalı parantezli varyantları temizle: (Mülga: ...)
        text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:[^)]*\)\s*', ' ', text, flags=re.IGNORECASE)
        
        # Açık parantezli varyantları temizle (satır sonuna kadar): (Mülga: ...
        text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:.*$', ' ', text, flags=re.IGNORECASE)
        
        # Yürürlük bilgilerini temizle
        text = re.sub(r'\s*\(?\s*Yürürlük\s*:.*$', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s*\b\d+\s*md\.\s*Yürürlük\s*:.*$', ' ', text, flags=re.IGNORECASE)
        
        # Fazla boşlukları temizle
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    @staticmethod
    def extract_clean_text(paragraph) -> Optional[str]:
        """
        Paragraftan temiz metni çıkarır.
        - TÜM run'ları toplar (strikethrough dahil)
        - Sonra sadece strikethrough olMAYANları birleştirir
        - Varyant bilgilerini temizler
        - Boş satırları atlar
        
        Returns:
            Temiz metin veya None (tamamen üstü çizili ise)
        """
        # Adım 1: TÜM run'ları topla
        all_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text.strip():
                continue
            
            # Her run'ı flag'leri ile birlikte topla
            all_parts.append({
                'text': text,
                'isStrikethrough': DocxStyleAnalyzer.is_strikethrough(run)
            })
        
        # Adım 2: Sadece strikethrough olmayanları filtrele
        clean_parts = [part['text'] for part in all_parts if not part['isStrikethrough']]
        
        if not clean_parts:
            return None
        
        combined_text = ''.join(clean_parts)
        
        # Adım 3: Varyant bilgilerini temizle
        cleaned_text = TextExtractor.clean_variants(combined_text)
        
        return cleaned_text.strip() if cleaned_text.strip() else None
    
    @staticmethod
    def should_skip_line(line: str) -> bool:
        """
        Bu satırı atlamalı mıyız?
        
        Atlanacak satırlar:
        - Boş satırlar
        - İçindekiler satırları (.... 25 formatı)
        - Sayfa numaraları (sadece rakam içeren satırlar)
        - "İÇİNDEKİLER" başlığı
        
        NOT: Varyant bilgileri (Mülga, Değişik, Ek) artık satır bazında atlanmıyor,
        clean_variants ile metin içinden temizleniyor.
        """
        if not line or not line.strip():
            return True
        
        line_stripped = line.strip()
        
        # İçindekiler bölümü - sadece başlık varsa
        if line_stripped.upper() == 'İÇİNDEKİLER':
            return True
        
        # İçindekiler formatı (.... 25)
        if RegexPatterns.TOC.search(line_stripped):
            return True
        
        # Sayfa numarası - sadece rakamlardan oluşan satır
        if RegexPatterns.PAGE_NUMBER.match(line_stripped):
            return True
        
        return False


class MetadataExtractor:
    """
    Chunk'lardan metadata çıkarır (SUT ekleri).
    RAG'de filtering için kullanılabilir.
    """
    
    @staticmethod
    def extract_sut_ekler(content: str) -> List[str]:
        """
        SUT ek referanslarını bulur.
        Örnek: "EK-2/A-2" -> ["EK-2/A-2"]
        """
        matches = RegexPatterns.SUT_EK.findall(content)
        normalized = [re.sub(r'\s+', '', m.upper()) for m in matches]
        return sorted(list(set(normalized)))


class SlidingWindowChunker:
    """Cümle bazlı sliding window ile chunking yapar."""
    
    @staticmethod
    def split_into_words(text: str) -> List[str]:
        """Metni kelimelere ayırır."""
        # Türkçe karakterleri de içeren kelime bölme
        words = re.findall(r'\S+', text)
        return words
    
    @staticmethod
    def find_sentence_end(words: List[str], start_index: int) -> int:
        """
        Verilen indeksten sonraki ilk cümle sonunu bulur.
        
        Args:
            words: Kelime listesi
            start_index: Aramaya başlanacak indeks 
        
        Returns:
            Cümle sonu indeksi (nokta ile biten kelime)
        """
        for i in range(start_index, len(words)):
            word = words[i]
            # Kelime nokta ile bitiyorsa
            if word.endswith('.'):
                return i + 1  # Bu kelimenin sonraki indeksini döndür (dahil etmek için)
        
        # Nokta bulunamazsa son kelimeyi döndür
        return len(words)
    
    @staticmethod
    def find_sentence_start(words: List[str], target_index: int) -> int:
        """
        Verilen indeksten geriye doğru en yakın cümle başlangıcını bulur.
        
        Args:
            words: Kelime listesi
            target_index: Hedef indeks (overlap için)
        
        Returns:
            Cümle başlangıç indeksi
        """
        # target_index'ten geriye doğru git
        for i in range(target_index - 1, -1, -1):
            word = words[i]
            # Bir önceki kelime nokta ile bitiyorsa, bu yeni cümlenin başlangıcı
            if word.endswith('.'):
                return i + 1  # Noktadan sonraki kelime
        
        # Hiç nokta bulunamazsa baştan başla
        return 0
    
    @staticmethod
    def create_chunks(words: List[str], chunk_size: int, overlap_size: int) -> List[str]:
        """
        Kelime listesinden cümle bazlı sliding window chunk'ları oluşturur.
        
        Mantık:
        1. Minimum chunk_size kelimeye ulaş (400)
        2. Sonraki ilk nokta (.) işaretine kadar devam et
        3. Overlap için geriye doğru overlap_size kelime civarındaki cümle başına git
        
        Args:
            words: Kelime listesi
            chunk_size: Minimum chunk kelime sayısı (300)
            overlap_size: Overlap kelime sayısı (100)
        
        Returns:
            Chunk'ların listesi (her biri string)
        """
        if not words:
            return []
        
        chunks = []
        start = 0
        
        while start < len(words):
            # Minimum chunk_size kelimeye ulaş
            min_end = start + chunk_size
            
            # Son chunk kontrolü - yeterli kelime kalmadıysa
            if min_end >= len(words):
                # Kalan tüm kelimeleri al
                chunk_words = words[start:]
                chunk_text = ' '.join(chunk_words)
                chunks.append(chunk_text)
                break
            
            # İlk noktayı bul (minimum 400. kelimeden sonra)
            actual_end = SlidingWindowChunker.find_sentence_end(words, min_end)
            
            # Chunk'ı al
            chunk_words = words[start:actual_end]
            chunk_text = ' '.join(chunk_words)
            chunks.append(chunk_text)
            
            # Bir sonraki chunk için başlangıç noktası
            # Overlap için geriye git
            overlap_target = actual_end - overlap_size
            
            # Overlap hedefinden geriye doğru en yakın cümle başını bul
            new_start = SlidingWindowChunker.find_sentence_start(words, max(overlap_target, 0))

            # Eğer ilerleme yoksa (sonsuz döngü önleme)
            if new_start <= start:
                start = actual_end
            else:
                start = new_start
        
        return chunks


class SutTebligDocx:
    """
    SUT Tebliği için sliding window chunker.
    300 kelime chunk + 100 kelime overlap
    """
    
    def __init__(self, chunk_size: int = 300, overlap_size: int = 100):
        self.chunk_size = chunk_size
        self.overlap_size = overlap_size

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        """
        DOCX'i işler ve chunk listesi döner.
        
        Returns:
            List[str] - Her biri chunk içeriği (string)
        """
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        
        # Adım 1: Temiz metni çıkar
        clean_lines = []
        content_started = False
        
        for para in self.doc.paragraphs:
            # Paragraftan temiz metni çıkar
            line = TextExtractor.extract_clean_text(para)
            
            if line is None:
                continue
            
            # İçerik başlamadıysa, ilk BÖLÜM başlığını bekle
            if not content_started:
                if RegexPatterns.BOLUM.match(line.strip()):
                    content_started = True
                else:
                    continue
            
            # Bu satırı atlamamız gerekiyor mu?
            if TextExtractor.should_skip_line(line):
                continue
            
            clean_lines.append(line)
        
        # Adım 2: Tüm satırları birleştir
        full_text = ' '.join(clean_lines)
        
        # Adım 3: Kelimelere ayır
        words = SlidingWindowChunker.split_into_words(full_text)
        
        # Adım 4: Sliding window ile chunk'la
        chunk_texts = SlidingWindowChunker.create_chunks(
            words, 
            self.chunk_size, 
            self.overlap_size
        )
        
        # Adım 5: Chunk'ları formatla ve metadata ekle
        chunks = []
        for chunk_text in chunk_texts:
            # SUT eklerini çıkar
            sut_ekler = MetadataExtractor.extract_sut_ekler(chunk_text)
            
            # Metadata'yı chunk sonuna ekle
            final_content = chunk_text
            if sut_ekler:
                final_content = f"{chunk_text} (SUT Ekleri: {', '.join(sut_ekler)})."
            
            chunks.append(final_content)
        
        return chunks


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Turkish", callback=None, **kwargs):
    """
    RAG Entegrasyon Fonksiyonu
    
    SUT Tebliği için sliding window chunking (300 kelime + 100 overlap)
    """
    
    parser_config = kwargs.get("parser_config", {})
    
    # Chunk size ve overlap'i parser_config'den al (yoksa default)
    chunk_size = parser_config.get("chunk_size", 300)
    overlap_size = parser_config.get("overlap_size", 100)
    
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    
    eng = lang.lower() == "english"
    
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "SUT Tebligi isleniyor (sliding window)...")
        rawChunks = SutTebligDocx(chunk_size, overlap_size)(filename, binary, from_page, to_page)
        callback(0.7, "SUT chunking tamamlandi.")
        
        result = []
        for chunkContent in rawChunks:
            d = copy.deepcopy(doc)
            d["content_with_weight"] = chunkContent
            d["content_ltks"] = rag_tokenizer.tokenize(chunkContent)
            d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
            result.append(d)
        
        callback(0.9, f"{len(result)} chunk olusturuldu (chunk_size: {chunk_size}, overlap: {overlap_size}).")
        return result
    
    else:
        raise NotImplementedError("SUT Teblig chunker sadece .docx dosyalarini destekler")


if __name__ == "__main__":
    import sys
    
    def dummy(prog=None, msg=""):
        pass
    
    chunk(sys.argv[1], callback=dummy)