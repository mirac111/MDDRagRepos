"""
HastaKabulModulu Chunker - RAGFlow Entegrasyonu
===============================================
BizMed HBYS Hasta Kabul modulu dokumantasyonu icin bold baslik bazli chunker.

Yapi:
  - Tamamen bold paragraf  →  yeni chunk baslar (baslik)
  - Normal paragraflar     →  o chunk'un icerigi

Chunk formati:
  {baslik}: {icerik paragraflari}
"""

import re
import copy
from docx import Document
from io import BytesIO

from rag.nlp import rag_tokenizer


# ── Yardimci fonksiyonlar ─────────────────────────────────────────────────────

def is_fully_bold(paragraph) -> bool:
    """Paragraftaki tum dolu run'lar bold ise True doner."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.bold for r in runs)


def get_heading_text(paragraph) -> str:
    """
    Bold paragrafin baslik metnini dondurur.
    run'lari direkt birlestirip fazla bosluklari temizler.
    """
    combined = ''.join(r.text for r in paragraph.runs)
    return re.sub(r'\s+', ' ', combined).strip()


def get_paragraph_text(paragraph) -> str:
    """Normal paragrafin temiz metnini dondurur."""
    text = paragraph.text.strip()
    return re.sub(r'\s+', ' ', text) if text else ''


# ── Ana chunker sinifi ────────────────────────────────────────────────────────

class HastaKabulDocx:
    """
    Hasta Kabul modulu dokumantasyonu icin bold baslik bazli chunker.
    Her baslik + altindaki paragraflar bir chunk olarak uretilir.
    """

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        doc = Document(filename) if not binary else Document(BytesIO(binary))

        chunks = []
        current_heading = None
        current_paragraphs = []

        def kaydet_chunk():
            if current_heading is None:
                return
            icerik = ' '.join(p for p in current_paragraphs if p)
            icerik = re.sub(r'\s+', ' ', icerik).strip()
            if icerik:
                chunk = f"{current_heading}: {icerik}"
            else:
                chunk = f"{current_heading}:"
            chunks.append(chunk)

        for para in doc.paragraphs:
            raw = para.text.strip()
            if not raw:
                continue

            if is_fully_bold(para):
                kaydet_chunk()
                current_heading = get_heading_text(para)
                current_paragraphs = []
            else:
                text = get_paragraph_text(para)
                if text and current_heading is not None:
                    current_paragraphs.append(text)

        kaydet_chunk()
        return chunks


# ── RAGFlow entegrasyon fonksiyonu ────────────────────────────────────────────

def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Turkish", callback=None, **kwargs):
    """
    RAGFlow Entegrasyon Fonksiyonu

    Hasta Kabul modulu dokumantasyonu icin bold baslik bazli chunking.
    Her baslik + icerigi ayri bir chunk olarak uretilir.
    """

    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Hasta Kabul dokumani isleniyor (bold baslik bazli chunking)...")
        rawChunks = HastaKabulDocx()(filename, binary, from_page, to_page)
        callback(0.7, "Chunking tamamlandi.")

        result = []
        for chunkContent in rawChunks:
            d = copy.deepcopy(doc)
            d["content_with_weight"] = chunkContent
            d["content_ltks"] = rag_tokenizer.tokenize(chunkContent)
            d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
            result.append(d)

        callback(0.9, f"{len(result)} chunk olusturuldu.")
        return result

    else:
        raise NotImplementedError("HastaKabul chunker sadece .docx dosyalarini destekler")


# ── Test modu ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], callback=dummy)