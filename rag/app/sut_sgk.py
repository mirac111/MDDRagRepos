import re
import copy
from docx import Document
from typing import List, Dict, Optional, Tuple
from io import BytesIO

from common.constants import ParserType
from rag.nlp import rag_tokenizer

# ── Ortak sabitler ────────────────────────────────────────────────────────────

TURKISH_ALPHA = ['a','b','c','ç','d','e','f','g','ğ','h','ı','i','j','k','l','m',
                 'n','o','ö','p','r','s','ş','t','u','ü','v','y','z']

TR_ALFABE = ['a','b','c','ç','d','e','f','g','ğ','h',
             'ı','i','j','k','l','m','n','o','ö','p',
             'r','s','ş','t','u','ü','v','y','z']

ROMAN = ['i','ii','iii','iv','v','vi','vii','viii','ix','x',
         'xi','xii','xiii','xiv','xv','xvi','xvii','xviii','xix','xx']

NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

MAX_TOKENS = 256

# ── Ortak LEVEL dict (en geniş — B4 bazlı) ───────────────────────────────────

LEVEL = {
    "BÖLÜM":          0,
    "MADDE_1":        1,
    "MADDE_2":        2,
    "MADDE_3":        3,
    "MADDE_4":        4,
    "MADDE_5":        5,
    "MADDE_6":        6,
    "FIKRA":          7,
    "BENT_HARF":      8,
    "BENT_NUM":       8,
    "BENT_HARF_NOKTA":8,
    "BENT_NUM_TIRE":  8,
    "METİN":          9,
}

# ── Ortak regexler ────────────────────────────────────────────────────────────

BOLUM_RE = re.compile(
    r"^(BİRİNCİ|İKİNCİ|ÜÇÜNCÜ|DÖRDÜNCÜ|BEŞİNCİ|ALTINCI|YEDİNCİ|SEKİZİNCİ|DOKUZUNCU|ONUNCU)\s+BÖLÜM$"
)
MADDE_RE_STRICT = re.compile(
    r"^(\d+\.\d+(?:\.\d+)*(?:\.[A-ZÇĞİÖŞÜa-zçğıöşü](?:\.\d+)?)?(?:-\d+)?)\s*[-–—]\s*(.+)$"
)
MADDE_RE_TOLERANT = re.compile(
    r"^(\d+\.\d+(?:\.\d+)*(?:\.[A-ZÇĞİÖŞÜa-zçğıöşü](?:\.\d+)?)?(?:-\d+)?)\s*[-–—\s]?\s*(.+)$"
)
FIKRA_RE     = re.compile(r"^\((\d+)\)\s*(.*)$")
BENT_HARF_RE = re.compile(
    r"^[\s\t]*\(?(?!(?:md|rg|vb|vs|bkz)\b)([a-zçğıöşü]{1,4})\s*([-\)\.])\s*(.*)$",
    re.IGNORECASE
)
BENT_NUM_RE      = re.compile(r"^(\d+)\)\s*(.*)$")
BENT_NUM_TIRE_RE = re.compile(r"^(\d+)\-\s*(.*)$")
TOC_RE           = re.compile(r"\.{4,}\s+\d+\s*$")

# ── Bölüm konfigürasyonları ───────────────────────────────────────────────────

BOLUMLER = [
    {
        "baslangic":   "BİRİNCİ BÖLÜM",
        "bitis":       "İKİNCİ BÖLÜM",
        "madde_max":   4,
        "madde_depth": "b1",
        "madde_re":    "strict",
        "bent_tipleri":("BENT_HARF", "BENT_NUM"),
        "in_degisik":  False,
        "bent_mantigi":"basit_b1",
        "roman_guard": False,
        "toc_b4":      False,
        "metin_filtre":True,
    },
    {
        "baslangic":   "İKİNCİ BÖLÜM",
        "bitis":       "ÜÇÜNCÜ BÖLÜM",
        "madde_max":   5,
        "madde_depth": "standart",
        "madde_re":    "strict",
        "bent_tipleri":("BENT_HARF", "BENT_NUM"),
        "in_degisik":  False,
        "bent_mantigi":"stack",
        "roman_guard": False,
        "toc_b4":      False,
        "metin_filtre":True,
    },
    {
        "baslangic":   "ÜÇÜNCÜ BÖLÜM",
        "bitis":       "DÖRDÜNCÜ BÖLÜM",
        "madde_max":   5,
        "madde_depth": "standart",
        "madde_re":    "tolerant",
        "bent_tipleri":("BENT_HARF", "BENT_NUM", "BENT_NUM_TIRE"),
        "in_degisik":  False,
        "bent_mantigi":"stack",
        "roman_guard": False,
        "toc_b4":      False,
        "metin_filtre":True,
    },
    {
        "baslangic":   "DÖRDÜNCÜ BÖLÜM",
        "bitis":       "BEŞİNCİ BÖLÜM",
        "max_tokens":  768, 
        "madde_max":   6,
        "madde_depth": "standart",
        "madde_re":    "tolerant",
        "bent_tipleri":("BENT_HARF", "BENT_NUM", "BENT_HARF_NOKTA", "BENT_NUM_TIRE"),
        "in_degisik":  True,
        "bent_mantigi":"stack",
        "roman_guard": False,
        "toc_b4":      True,
        "metin_filtre":True,
    },
    {
        "baslangic":   "BEŞİNCİ BÖLÜM",
        "bitis":       "ALTINCI BÖLÜM",
        "madde_max":   5,
        "madde_depth": "standart",
        "madde_re":    "tolerant",
        "bent_tipleri":("BENT_HARF", "BENT_NUM"),
        "in_degisik":  False,
        "bent_mantigi":"stack",
        "roman_guard": True,
        "toc_b4":      False,
        "metin_filtre":True,
    },
    {
        "baslangic":   "ALTINCI BÖLÜM",
        "bitis":       "YEDİNCİ BÖLÜM",
        "madde_max":   5,
        "madde_depth": "standart",
        "madde_re":    "tolerant",
        "bent_tipleri":("BENT_HARF", "BENT_NUM"),
        "in_degisik":  False,
        "bent_mantigi":"basit_b6",
        "roman_guard": False,
        "toc_b4":      False,
        "metin_filtre":False,
    },
]

# ── Madde derinlik fonksiyonları ──────────────────────────────────────────────

def madde_depth_b1(text: str) -> int:
    m = re.match(r'^([\d\.]+[A-ZÇĞİÖŞÜa-zçğıöşü]?(?:-\d+)?)', text.strip())
    if not m:
        return 0
    num = m.group(1)
    base_parts = num.split('.')
    depth = len(base_parts) - 2
    if '-' in base_parts[-1]:
        depth += 1
    return max(0, depth)

def madde_depth_standart(text: str) -> int:
    t = re.sub(r'[–—‐]', '-', text.strip())
    t = re.sub(r'\.([A-ZÇĞİÖŞÜa-zçğıöşü]+)\.(\d+)', r'.\1-\2', t)
    m = re.match(r'^(\d+(?:\.\d+)+)(?:\.([A-ZÇĞİÖŞÜa-zçğıöşü]+))?((?:-[A-ZÇĞİÖŞÜa-zçğıöşü\d\.]+)+)?', t)
    if not m:
        return 0
    base     = m.group(1)
    harf     = m.group(2)
    sonekler = m.group(3)
    depth = len(base.split('.')) - 2
    if harf:
        depth += 1
    if sonekler:
        parcalar = re.findall(r'-([A-ZÇĞİÖŞÜa-zçğıöşü]{1,2}|\d+(?:\.\d+)*)(?=-|$)', sonekler)
        for p in parcalar:
            depth += 1 + p.count('.')
    return max(0, depth)

# ── Yardımcı fonksiyonlar ─────────────────────────────────────────────────────

def normalize_dashes(text: str) -> str:
    return re.sub(r'[–—‐]', '-', text)

def normalize_madde_format(text: str) -> str:
    while re.search(r'(\d+)\.\s+(\d)', text):
        text = re.sub(r'(\d+)\.\s+(\d)', r'\1.\2', text)
    text = re.sub(r'\.([A-ZÇĞİÖŞÜ])\.\s+', r'.\1-', text)
    text = re.sub(r'(\d+)\.\s+([A-ZÇĞİÖŞÜ])', r'\1-\2', text)
    text = re.sub(r'^(\d+\.\d+(?:\.\d+)*(?:\.[A-ZÇĞİÖŞÜ])?)\s+(\()', r'\1 - \2', text)
    return text

def clean_variants(text: str) -> str:
    if not text:
        return text
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:[^)]*\)\s*', ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:.*$',        ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\(?\s*Yürürlük\s*:.*$',                  ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s*\b\d+\s*md\.\s*Yürürlük\s*:.*$',        ' ', text, flags=re.IGNORECASE)
    return re.sub(r'\s+', ' ', text).strip()

def is_strikethrough(run) -> bool:
    try:
        return bool(run.font.strike or run.font.double_strike)
    except:
        return False

def is_fully_strikethrough(segments: List[Dict]) -> bool:
    if not segments:
        return False
    return all(seg['isStrikethrough'] for seg in segments)

def is_toc_line(line: str, toc_b4: bool = False) -> bool:
    if bool(TOC_RE.search(line)) or 'İÇİNDEKİLER' in line.upper():
        return True
    if re.match(r'^\d{1,4}$', line.strip()):
        return True
    if toc_b4:
        if re.match(r'^[A-ZÇĞİÖŞÜ]?\d{1,4}[A-ZÇĞİÖŞÜa-zçğıöşü]?\d*$', line.strip()):
            return True
    else:
        if re.search(r'[a-zçğıöşüA-ZÇĞİÖŞÜ]\d+$', line):
            return True
    return False

# ── Tip tespiti ───────────────────────────────────────────────────────────────

def classify(clean_text: str, cfg: dict) -> str:
    clean_text = clean_text.strip()
    t = normalize_dashes(clean_text)
    t = normalize_madde_format(t)

    if BOLUM_RE.match(clean_text):
        return "BÖLÜM"
    if re.match(r'^\d+\.\s*\d+(?:\.\d+)*(?:\.[A-ZÇĞİÖŞÜa-zçğıöşü])?(?:-\d+)?\s*,', t):
        return "METİN"

    madde_re = MADDE_RE_STRICT if cfg["madde_re"] == "strict" else MADDE_RE_TOLERANT
    if madde_re.match(t):
        depth_fn = madde_depth_b1 if cfg["madde_depth"] == "b1" else madde_depth_standart
        d = depth_fn(t)
        level = min(d + 1, cfg["madde_max"])
        return f"MADDE_{level}"

    if FIKRA_RE.match(clean_text):
        return "FIKRA"
    if BENT_NUM_RE.match(clean_text):
        return "BENT_NUM"
    if "BENT_NUM_TIRE" in cfg["bent_tipleri"] and BENT_NUM_TIRE_RE.match(clean_text):
        return "BENT_NUM_TIRE"
    harf_m = BENT_HARF_RE.match(clean_text)
    if harf_m:
        if "BENT_HARF_NOKTA" in cfg["bent_tipleri"] and harf_m.group(2) == '.':
            return "BENT_HARF_NOKTA"
        return "BENT_HARF"
    return "METİN"

# ── Segment çıkarma ───────────────────────────────────────────────────────────

def extract_segments(para, document, list_counters: dict,
                     in_degisik_mod: bool = False) -> Tuple[List[Dict], int]:
    segments  = []
    para_ilvl = 0
    word_num_text = ""
    try:
        if (para._element.pPr is not None and
                para._element.pPr.numPr is not None and
                para._element.pPr.numPr.numId is not None):
            numPr = para._element.pPr.numPr
            numId = numPr.numId.val
            ilvl  = numPr.ilvl.val if numPr.ilvl is not None else 0
            para_ilvl = ilvl
            key   = (numId, ilvl)
            if key not in list_counters:
                list_counters[key] = 0
            list_counters[key] += 1
            count = list_counters[key]
            fmt, _, start = _get_numbering_format(para, document)
            if fmt == 'lowerLetter':
                LATIN = 'abcdefghijklmnopqrstuvwxyz'
                effective = count + start - 1
                idx    = (effective - 1) % 26
                repeat = (effective - 1) // 26 + 1
                letter = LATIN[idx] * repeat
                word_num_text = f"{letter}) "
            elif fmt == 'upperLetter':
                effective = count + start - 1
                letter = TURKISH_ALPHA[effective - 1].upper() if effective <= len(TURKISH_ALPHA) else str(effective)
                word_num_text = f"{letter}) "
            else:
                effective = count + start - 1
                word_num_text = f"{effective}) "
    except:
        pass

    if in_degisik_mod:
        in_degisik = False
        for run in para.runs:
            text = run.text
            if not text.strip():
                continue
            if re.search(r'\(?\s*(Mülga|Değişik|Ek)\s*:', text, re.IGNORECASE):
                in_degisik = True
            if in_degisik:
                segments.append({'text': text, 'isStrikethrough': True})
                if ')' in text:
                    in_degisik = False
            else:
                segments.append({'text': text, 'isStrikethrough': is_strikethrough(run)})
    else:
        for run in para.runs:
            text = run.text
            if not text.strip():
                continue
            segments.append({'text': text, 'isStrikethrough': is_strikethrough(run)})

    if word_num_text and segments:
        segments.insert(0, {
            'text': word_num_text,
            'isStrikethrough': segments[0]['isStrikethrough']
        })
    return segments, para_ilvl

def _get_numbering_format(paragraph, document) -> Tuple[Optional[str], int, int]:
    try:
        numPr = paragraph._element.pPr.numPr
        if numPr is None:
            return None, 0, 1
        numId = numPr.numId.val
        ilvl  = numPr.ilvl.val if numPr.ilvl is not None else 0
        numberingPart = document.part.numbering_part
        if numberingPart is None:
            return None, ilvl, 1
        numbering = numberingPart.element
        for num in numbering.findall(f'.//{NS}num'):
            if num.get(f'{NS}numId') and int(num.get(f'{NS}numId')) == numId:
                absIdEl = num.find(f'.//{NS}abstractNumId')
                if absIdEl is None:
                    continue
                absVal = absIdEl.get(f'{NS}val')
                for absNum in numbering.findall(f'.//{NS}abstractNum'):
                    if absNum.get(f'{NS}abstractNumId') == absVal:
                        for lvl in absNum.findall(f'.//{NS}lvl'):
                            if lvl.get(f'{NS}ilvl') and int(lvl.get(f'{NS}ilvl')) == ilvl:
                                fmt      = lvl.find(f'.//{NS}numFmt')
                                start_el = lvl.find(f'.//{NS}start')
                                start_val = int(start_el.get(f'{NS}val', '1')) if start_el is not None else 1
                                if fmt is not None:
                                    return fmt.get(f'{NS}val'), ilvl, start_val
        return None, ilvl, 1
    except:
        return None, 0, 1

# ── Girinti hesapla ───────────────────────────────────────────────────────────

def get_indent(tip: str, bonus: int = 0, last_madde_depth: int = 0) -> str:
    if tip.startswith("MADDE_"):
        total = int(tip.split("_")[1]) - 1
    elif tip == "FIKRA":
        total = last_madde_depth + 1
    elif tip in ("BENT_HARF", "BENT_NUM", "BENT_HARF_NOKTA", "BENT_NUM_TIRE"):
        total = last_madde_depth + 2 + bonus
    elif tip == "METİN":
        total = last_madde_depth + 1
    else:
        total = 0
    return "    " * total

# ── Bent stack mantığı (B2-5) ─────────────────────────────────────────────────

def bent_stack_guncelle(tip: str, clean_text: str, bent_stack: list,
                        roman_guard: bool) -> list:
    harf_m = re.match(r'^\(?([a-zçğıöşü]+)[)\.\-]', clean_text.strip(), re.IGNORECASE)
    num_m  = re.match(r'^(\d+)[)\-]', clean_text.strip())
    deger  = harf_m.group(1) if harf_m else (num_m.group(1) if num_m else "")

    stack_tipler = [s[0] for s in bent_stack]
    if tip not in stack_tipler:
        return bent_stack + [(tip, deger)]

    onceki_deger = ""
    for i in range(len(bent_stack)-1, -1, -1):
        if bent_stack[i][0] == tip:
            onceki_deger = bent_stack[i][1]
            break

    try:
        is_kardes = int(deger) > int(onceki_deger)
    except:
        is_kardes_tr = False
        is_kardes_roman = False
        if deger and onceki_deger:
            def tr_sira(s):
                return (len(s), TR_ALFABE.index(s[0]) if s[0] in TR_ALFABE else -1)
            is_kardes_tr = tr_sira(deger) > tr_sira(onceki_deger)
        if deger in ROMAN and onceki_deger in ROMAN:
            is_kardes_roman = ROMAN.index(deger) > ROMAN.index(onceki_deger)
        is_kardes = is_kardes_tr or is_kardes_roman

    if roman_guard and is_kardes and tip == "BENT_HARF" and deger in ROMAN:
        if any(s[0] == "BENT_NUM" for s in bent_stack):
            is_kardes = False
            if bent_stack and bent_stack[-1][0] == "BENT_HARF" and bent_stack[-1][1] in ROMAN:
                return bent_stack[:-1] + [(tip, deger)]

    if not is_kardes:
        return bent_stack + [(tip, deger)]

    best_pos = None
    for idx_s, s in enumerate(bent_stack):
        if s[0] != tip:
            continue
        prev_val = s[1]
        try:
            if int(deger) == int(prev_val) + 1:
                best_pos = idx_s
        except:
            def tr_next(s, p):
                if len(s) == len(p):
                    if s[0] in TR_ALFABE and p[0] in TR_ALFABE:
                        return TR_ALFABE.index(s[0]) == TR_ALFABE.index(p[0]) + 1
                elif len(s) == len(p) + 1:
                    return all(c == TR_ALFABE[-1] for c in p) and all(c == TR_ALFABE[0] for c in s)
                return False
            if tr_next(deger, prev_val):
                best_pos = idx_s
            if deger in ROMAN and prev_val in ROMAN:
                if ROMAN.index(deger) == ROMAN.index(prev_val) + 1:
                    best_pos = idx_s

    if best_pos is None:
        best_pos = next(
            (i for i, s in reversed(list(enumerate(bent_stack))) if s[0] == tip),
            None
        )

    if best_pos is not None:
        new_stack = bent_stack[:best_pos + 1]
        new_stack[-1] = (tip, deger)
        return new_stack
    else:
        return bent_stack + [(tip, deger)]

# ── Chunk üretim fonksiyonları ────────────────────────────────────────────────

def token_say(text: str) -> int:
    return len(text.split())

def chunk_olustur(path_titles: list, lines: list) -> str:
    tum_satirlar = [s for s in path_titles + lines if s.strip()]
    return "\n".join(tum_satirlar)

def greedy_grupla(sonuclar: list, max_tokens: int) -> list:
    gruplar = []
    mevcut_token = 0
    mevcut_lines = []
    for sonuc in sonuclar:
        if sonuc["yazildi"]:
            if mevcut_lines:
                gruplar.append({"token": mevcut_token, "lines": mevcut_lines})
                mevcut_token = 0
                mevcut_lines = []
            continue
        c_token = sonuc["token"]
        c_lines = sonuc["lines"]
        if mevcut_lines and (mevcut_token + c_token > max_tokens):
            gruplar.append({"token": mevcut_token, "lines": mevcut_lines})
            mevcut_token = c_token
            mevcut_lines = c_lines[:]
        else:
            mevcut_token += c_token
            mevcut_lines.extend(c_lines)
    if mevcut_lines:
        gruplar.append({"token": mevcut_token, "lines": mevcut_lines})
    return gruplar

def process(node: dict, path_titles: list, all_chunks: list, max_tokens: int) -> dict:
    kendi_token = token_say(node["text"])
    kendi_line  = node.get("indent", "") + node["text"]
    if not node.get("children"):
        return {"token": kendi_token, "lines": [kendi_line], "yazildi": False}
    cocuk_path = path_titles + [node["text"]]
    cocuk_sonuclari = []
    for cocuk in node["children"]:
        cocuk_sonuclari.append(process(cocuk, cocuk_path, all_chunks, max_tokens))
    gruplar = greedy_grupla(cocuk_sonuclari, max_tokens)
    yazilmamis_token = sum(s["token"] for s in cocuk_sonuclari if not s["yazildi"])
    toplam_token = kendi_token + yazilmamis_token
    if len(gruplar) == 1 and toplam_token <= max_tokens:
        return {
            "token":   toplam_token,
            "lines":   [kendi_line] + gruplar[0]["lines"],
            "yazildi": False
        }
    for grup in gruplar:
        all_chunks.append(chunk_olustur(cocuk_path, grup["lines"]))
    return {"token": kendi_token, "lines": [kendi_line], "yazildi": True}

def json_to_chunks(tree: list, max_tokens: int) -> list:
    all_chunks = []
    for top_node in tree:
        sonuc = process(top_node, [], all_chunks, max_tokens)
        if not sonuc["yazildi"]:
            all_chunks.append(chunk_olustur([], sonuc["lines"]))
    return all_chunks

# ── Tek bölümü işle ───────────────────────────────────────────────────────────

def isle_bolum(doc, cfg: dict) -> list:
    max_tokens = cfg.get("max_tokens", MAX_TOKENS)
    in_content       = False
    skip_metin_count = 0
    list_counters    = {}
    indent_bonus     = 0
    prev_bent        = ""
    prev_harf_len    = 0
    last_madde_depth = 0
    prev_indent      = ""
    bent_stack       = []

    root  = {"tip": "ROOT", "text": "", "children": [], "_level": -1}
    stack = [root]

    baslangic = cfg["baslangic"]
    bitis     = cfg["bitis"]

    for para in doc.paragraphs:
        segments, para_ilvl = extract_segments(
            para, doc, list_counters,
            in_degisik_mod=cfg["in_degisik"]
        )
        if not segments:
            continue

        raw_text = ''.join(seg['text'] for seg in segments).strip()
        if not raw_text:
            continue

        if not in_content:
            if raw_text == baslangic:
                in_content = True
            continue

        if bitis and raw_text == bitis:
            break

        if is_toc_line(raw_text, toc_b4=cfg["toc_b4"]):
            continue

        if is_fully_strikethrough(segments):
            continue

        clean_text = ''.join(
            seg['text'] for seg in segments if not seg['isStrikethrough']
        ).strip()
        if not clean_text:
            continue

        clean_text = clean_variants(clean_text)
        if not clean_text:
            continue

        tip = classify(clean_text, cfg)

        if tip == "BÖLÜM":
            continue

        if tip == "METİN" and skip_metin_count < 1:
            skip_metin_count += 1
            continue

        if tip.startswith("MADDE_"):
            level = int(tip.split("_")[1])
            last_madde_depth = level - 1
            indent_bonus  = 0
            prev_bent     = ""
            prev_harf_len = 0
            bent_stack    = []

        elif tip == "FIKRA":
            bent_stack    = []
            indent_bonus  = 0
            prev_bent     = ""
            prev_harf_len = 0

        elif tip in cfg["bent_tipleri"]:
            mantik = cfg["bent_mantigi"]

            if mantik == "stack":
                bent_stack   = bent_stack_guncelle(tip, clean_text, bent_stack, cfg["roman_guard"])
                indent_bonus = len(bent_stack) - 1

            elif mantik == "basit_b1":
                if prev_bent == "BENT_HARF" and tip == "BENT_NUM":
                    if prev_harf_len > 1:
                        indent_bonus = max(0, indent_bonus - 1)
                    else:
                        indent_bonus += 1
                elif prev_bent == "BENT_NUM" and tip == "BENT_HARF":
                    harf_m = re.match(r'^([a-zçğıöşü]+)\)', clean_text.strip(), re.IGNORECASE)
                    harf_len = len(harf_m.group(1)) if harf_m else 1
                    if harf_len > 1:
                        indent_bonus += 1
                    else:
                        indent_bonus = 0

            elif mantik == "basit_b6":
                if prev_bent == "BENT_HARF" and tip == "BENT_NUM":
                    indent_bonus += 1
                elif prev_bent == "BENT_NUM" and tip == "BENT_HARF":
                    indent_bonus = 0

            prev_bent = tip
            harf_m2 = re.match(r'^([a-zçğıöşü]+)\)', clean_text.strip(), re.IGNORECASE)
            prev_harf_len = len(harf_m2.group(1)) if harf_m2 else 0

        else:
            indent_bonus  = 0
            prev_bent     = ""
            prev_harf_len = 0

        indent = get_indent(tip, indent_bonus, last_madde_depth)

        if cfg["metin_filtre"]:
            if tip == "METİN" and len(clean_text.strip()) == 1:
                continue
            if tip == "METİN" and clean_text.strip().lower() in ("veya",):
                continue
            if tip == "METİN":
                indent = prev_indent

        prev_indent = indent

        current_level = LEVEL.get(tip, 9)
        if tip in cfg["bent_tipleri"]:
            current_level = current_level + indent_bonus

        while len(stack) > 1 and (
            stack[-1]["_level"] > current_level or
            (stack[-1]["_level"] == current_level and stack[-1]["tip"] in (
                "MADDE_1", "MADDE_2", "MADDE_3", "MADDE_4", "MADDE_5", "MADDE_6",
                "BENT_HARF", "BENT_NUM", "BENT_HARF_NOKTA", "BENT_NUM_TIRE",
                "FIKRA", "METİN"
            ))
        ):
            stack.pop()

        node = {
            "tip":      tip,
            "text":     clean_text,
            "indent":   indent,
            "children": [],
            "_level":   current_level
        }
        stack[-1]["children"].append(node)
        stack.append(node)

    return json_to_chunks(root["children"], max_tokens)

# ── RAGFlow Entegrasyon Fonksiyonu ────────────────────────────────────────────

def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Turkish", callback=None, **kwargs):
    """
    RAGFlow Standart Chunking Arayüzü
    SGK Saglik Uygulama Tebligi icin DFS+Greedy chunking.
    """
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "SGK SUT Tebligi isleniyor (DFS+Greedy)...")

        docx_doc = Document(filename) if not binary else Document(BytesIO(binary))

        raw_chunks = []
        for cfg in BOLUMLER:
            raw_chunks.extend(isle_bolum(docx_doc, cfg))

        callback(0.7, "Chunking tamamlandi, tokenize ediliyor...")

        result = []
        for chunk_content in raw_chunks:
            d = copy.deepcopy(doc)
            d["content_with_weight"] = chunk_content
            d["content_ltks"]        = rag_tokenizer.tokenize(chunk_content)
            d["content_sm_ltks"]     = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
            result.append(d)

        callback(0.9, f"{len(result)} chunk olusturuldu.")
        return result

    else:
        raise NotImplementedError("SUT SGK chunker sadece .docx dosyalarini destekler.")


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Kullanim: python sut_sgk.py <dosya.docx>")
        sys.exit(1)

    doc = Document(sys.argv[1])
    tum_chunks = []
    for cfg in BOLUMLER:
        chunks = isle_bolum(doc, cfg)
        tum_chunks.extend(chunks)
        print(f"Bolum ({cfg['baslangic']}) -> {len(chunks)} chunk")

    print(f"\nToplam {len(tum_chunks)} chunk uretildi.")