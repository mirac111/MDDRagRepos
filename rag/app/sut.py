import re
import copy
from docx import Document
from typing import List, Dict, Optional
from io import BytesIO

from api.db import ParserType
from rag.nlp import rag_tokenizer


class RegexPatterns:
    BOLUM = re.compile(
        r"^(BİRİNCİ|İKİNCİ|ÜÇÜNCÜ|DÖRDÜNCÜ|BEŞİNCİ|ALTINCI|YEDİNCİ|SEKİZİNCİ|DOKUZUNCU|ONUNCU)\s+BÖLÜM$"
    )
    
    MADDE = re.compile(
        r"^(\d+\.\d+(?:\.\d+)*(?:\.[A-ZÇĞİÖŞÜa-zçğıöşü](?:\.\d+)?)?(?:-\d+)?)\s*[-–—]\s*(.+)$"
    )
    
    FIKRA = re.compile(r"^\((\d+)\)\s*(.*)$")
    
    BENT_HARF = re.compile(
        r"^\(?(?!(?:md|rg|vb|vs|bkz)\b)([a-zçğıöşü]{1,3})\s*([\)\.])\s*(.*)$", 
        re.IGNORECASE
    )
    
    BENT_NUMERIK = re.compile(r"^(\d+)\)\s*(.*)$")
    
    ALT_BENT = re.compile(
        r"^([a-zçğıöşü]-|\d+-?)\s*([\)\.]?)\s*(.*)$", 
        re.IGNORECASE
    )
    
    TOC = re.compile(r"\.{4,}\s+\d+\s*$")
    
    SUT_EK = re.compile(
        r'EK-\s*[1-4]\s*/\s*[A-ZÇĞİÖŞÜ](?:\s*-\s*\d+)?', 
        re.IGNORECASE
    )
    
    LAW_NUMBER = re.compile(r'\b(\d{3,5})\s{1,3}sayılı', re.IGNORECASE)
    
    PRICE_FORMATTING = re.compile(r'[\.…,\s]{3,}\s*(\d+(?:[.,]\d+)?)\s*TL')


class TextCleaner:
    @staticmethod
    def clean_variants(text: str) -> str:
        if not text:
            return text
        
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:[^)]*\)\s*', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s*\(\s*(Mülga|Değişik|Ek)\s*:.*$', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s*\(?\s*Yürürlük\s*:.*$', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s*\b\d+\s*md\.\s*Yürürlük\s*:.*$', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    @staticmethod
    def clean_price_formatting(text: str) -> str:
        if not text:
            return text
        return RegexPatterns.PRICE_FORMATTING.sub(r' : \1 TL', text)
    
    @staticmethod
    def normalize_dashes(text: str) -> str:
        return re.sub(r'[–—‐]', '-', text)
    
    @staticmethod
    def normalize_madde_format(text: str) -> str:
        while re.search(r'(\d+)\.\s+(\d)', text):
            text = re.sub(r'(\d+)\.\s+(\d)', r'\1.\2', text)
        
        text = re.sub(r'\.([A-ZÇĞİÖŞÜ])\.\s+', r'.\1-', text)
        text = re.sub(r'(\d+)\.\s+([A-ZÇĞİÖŞÜ])', r'\1-\2', text)
        text = re.sub(r'^(\d+\.\d+(?:\.\d+)*(?:\.[A-ZÇĞİÖŞÜ])?)\s+(\()', r'\1 - \2', text)
        
        return text


class MetadataExtractor:
    @staticmethod
    def extract_sut_ekler(content: str) -> List[str]:
        matches = RegexPatterns.SUT_EK.findall(content)
        normalized = [re.sub(r'\s+', '', m.upper()) for m in matches]
        return sorted(list(set(normalized)))
    
    @staticmethod
    def extract_law_numbers(content: str) -> List[str]:
        matches = RegexPatterns.LAW_NUMBER.findall(content)
        return sorted(list(set(matches)))


class DocxStyleAnalyzer:
    @staticmethod
    def is_strikethrough(run) -> bool:
        """✅ FIX: None değil bool döndür"""
        try:
            strike = run.font.strike
            double_strike = run.font.double_strike
            # None or None = None olmasını önle
            return bool(strike) or bool(double_strike)
        except:
            return False
    
    @staticmethod
    def is_red(run) -> bool:
        """✅ FIX: Her zaman bool döndür"""
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                r, g, b = rgb[0] if len(rgb) > 0 else 0, rgb[1] if len(rgb) > 1 else 0, rgb[2] if len(rgb) > 2 else 0
                return bool(r > 150 and g < 100 and b < 100)
            return False
        except:
            return False
    
    @staticmethod
    def has_numbering(paragraph) -> bool:
        """✅ FIX: Her zaman bool döndür"""
        try:
            if paragraph._element.pPr is None:
                return False
            numPr = paragraph._element.pPr.numPr
            if numPr is None:
                return False
            if numPr.numId is None:
                return False
            # numId.val None olabilir veya 0 olabilir
            val = numPr.numId.val
            return bool(val) if val is not None else False
        except:
            return False
    
    @staticmethod
    def get_numbering_format(paragraph, document) -> Optional[str]:
        try:
            numPr = paragraph._element.pPr.numPr
            if numPr is None:
                return None
            
            numId = numPr.numId.val
            ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
            
            numberingPart = document.part.numbering_part
            if numberingPart is None:
                return None
            
            numbering = numberingPart.element
            
            for num in numbering.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'):
                numIdAttr = num.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                if numIdAttr and int(numIdAttr) == numId:
                    abstractNumId = num.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                    if abstractNumId is not None:
                        absNumIdVal = abstractNumId.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        
                        for absNum in numbering.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum'):
                            absNumIdAttr = absNum.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                            if absNumIdAttr and absNumIdAttr == absNumIdVal:
                                for lvl in absNum.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lvl'):
                                    lvlVal = lvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                                    if lvlVal and int(lvlVal) == ilvl:
                                        numFmt = lvl.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt')
                                        if numFmt is not None:
                                            return numFmt.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            return None
        except:
            return None


class TurkishLetterConverter:
    ALPHABET = ['a', 'b', 'c', 'ç', 'd', 'e', 'f', 'g', 'ğ', 'h', 'ı', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'ö', 'p', 'r', 's', 'ş', 't', 'u', 'ü', 'v', 'y', 'z']
    
    @classmethod
    def number_to_letter(cls, num: int, is_upper: bool = False) -> str:
        if num < 1 or num > len(cls.ALPHABET):
            return str(num)
        
        letter = cls.ALPHABET[num - 1]
        return letter.upper() if is_upper else letter


class ParagraphSegmentExtractor:
    """✅ TAM DOĞRU: ragChunkingTest.py'den"""
    
    @staticmethod
    def extract(paragraph, listCounters: dict, document) -> List[Dict]:
        segments = []
        
        wordNumberText = ""
        hasWordNumbering = False
        
        if DocxStyleAnalyzer.has_numbering(paragraph):
            try:
                numPr = paragraph._element.pPr.numPr
                numId = numPr.numId.val
                ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
                
                key = (numId, ilvl)
                if key not in listCounters:
                    listCounters[key] = 0
                listCounters[key] += 1
                
                numFormat = DocxStyleAnalyzer.get_numbering_format(paragraph, document)
                
                if numFormat == 'lowerLetter':
                    wordNumberText = f"{TurkishLetterConverter.number_to_letter(listCounters[key], False)})"
                elif numFormat == 'upperLetter':
                    wordNumberText = f"{TurkishLetterConverter.number_to_letter(listCounters[key], True)})"
                else:
                    wordNumberText = f"{listCounters[key]})"
                
                hasWordNumbering = True
            except:
                pass
        
        for run in paragraph.runs:
            text = run.text
            if not text.strip():
                continue
            
            segments.append({
                'text': text,
                'isStrikethrough': DocxStyleAnalyzer.is_strikethrough(run),
                'isRed': DocxStyleAnalyzer.is_red(run)
            })
        
        # ✅ KRİTİK: Word numbering'i AYRI segment olarak INSERT et
        if hasWordNumbering and segments:
            firstRunStrikethrough = segments[0]['isStrikethrough'] if segments else False
            
            segments.insert(0, {
                'text': wordNumberText + '\t',
                'isStrikethrough': firstRunStrikethrough,
                'isRed': False
            })
        
        # ✅ KRİTİK: TÜM segment'leri döndür (filtreleme YOK!)
        return segments
    
    @staticmethod
    def merge_consecutive(segments: List[Dict]) -> List[Dict]:
        if not segments:
            return []
        
        merged = []
        current = segments[0].copy()
        
        for i in range(1, len(segments)):
            nextSeg = segments[i]
            
            if (nextSeg['isStrikethrough'] == current['isStrikethrough'] and 
                nextSeg['isRed'] == current['isRed']):
                current['text'] += nextSeg['text']
            else:
                merged.append(current)
                current = nextSeg.copy()
        
        merged.append(current)
        return merged


class HeaderDetector:
    """✅ TAM DOĞRU: ragChunkingTest.py'den"""
    
    @staticmethod
    def is_toc_line(line: str) -> bool:
        """✅ FIX: Her zaman bool döndür"""
        try:
            return bool(RegexPatterns.TOC.search(line)) or bool('İÇİNDEKİLER' in line.upper())
        except:
            return False
    
    @staticmethod
    def is_fully_strikethrough(segments: List[Dict]) -> bool:
        """✅ FIX: Her zaman bool döndür"""
        try:
            if not segments:
                return False
            return all(bool(seg.get('isStrikethrough', False)) for seg in segments)
        except:
            return False
    
    @staticmethod
    def detect(lineText: str, segments: List[Dict]) -> Optional[Dict]:
        lineText = lineText.strip()
        
        # ✅ TOC kontrolü
        if not lineText or HeaderDetector.is_toc_line(lineText):
            return None
        
        # ✅ Tam strikethrough kontrolü
        if HeaderDetector.is_fully_strikethrough(segments):
            return None
        
        # ✅ Strikethrough olmayan segment'leri filtrele
        filteredSegments = [seg for seg in segments if not seg.get('isStrikethrough', False)]
        if not filteredSegments:
            return None
        
        cleanLineText = ''.join([seg['text'] for seg in filteredSegments]).strip()
        
        if RegexPatterns.BOLUM.match(cleanLineText):
            return {
                'type': 'BOLUM',
                'text': cleanLineText,
                'level': 0
            }
        
        # ✅ Normalizasyon
        lineNormalized = TextCleaner.normalize_dashes(cleanLineText)
        lineNormalized = TextCleaner.normalize_madde_format(lineNormalized)
        
        # ✅ Virgüllü madde numarası kontrolü (liste değil başlık)
        if re.match(r'^\d+\.\s*\d+(?:\.\d+)*(?:\.[A-ZÇĞİÖŞÜa-zçğıöşü])?(?:-\d+)?\s*,', lineNormalized):
            return None
        
        maddeMatch = RegexPatterns.MADDE.match(lineNormalized)
        if maddeMatch:
            maddeTitle = TextCleaner.clean_variants(maddeMatch.group(2).strip())
            displayText = maddeTitle[:100] + '...' if len(maddeTitle) > 100 else maddeTitle
            
            return {
                'type': 'MADDE',
                'number': maddeMatch.group(1),
                'title': displayText,
                'fullTitle': maddeTitle,  # ✅ Tam başlık
                'level': 1
            }
        
        fikraMatch = RegexPatterns.FIKRA.match(cleanLineText)
        if fikraMatch:
            fikraContent = TextCleaner.clean_variants(fikraMatch.group(2).strip())
            displayText = fikraContent[:100] + '...' if len(fikraContent) > 100 else fikraContent
            
            return {
                'type': 'FIKRA',
                'number': fikraMatch.group(1),
                'content': displayText,
                'fullContent': fikraContent,  # ✅ Tam içerik
                'level': 2
            }
        
        # ✅ Önce numerik bent kontrol et
        numerikMatch = RegexPatterns.BENT_NUMERIK.match(cleanLineText)
        if numerikMatch:
            bentContent = TextCleaner.clean_variants(numerikMatch.group(2).strip())
            displayText = bentContent[:100] + '...' if len(bentContent) > 100 else bentContent
            
            return {
                'type': 'BENT',
                'subtype': 'NUMERIK',  # ✅ Alt tip
                'number': numerikMatch.group(1),
                'format': ')',
                'content': displayText,
                'fullContent': bentContent,  # ✅ Tam içerik
                'level': 3
            }
        
        bentMatch = RegexPatterns.BENT_HARF.match(cleanLineText)
        if bentMatch:
            bentContent = TextCleaner.clean_variants(bentMatch.group(3).strip())
            displayText = bentContent[:100] + '...' if len(bentContent) > 100 else bentContent
            
            return {
                'type': 'BENT',
                'subtype': 'HARF',  # ✅ Alt tip
                'number': bentMatch.group(1),
                'format': bentMatch.group(2),
                'content': displayText,
                'fullContent': bentContent,  # ✅ Tam içerik
                'level': 3
            }
        
        altBentMatch = RegexPatterns.ALT_BENT.match(cleanLineText)
        if altBentMatch:
            altBentContent = TextCleaner.clean_variants(altBentMatch.group(3).strip())
            displayText = altBentContent[:100] + '...' if len(altBentContent) > 100 else altBentContent
            
            return {
                'type': 'ALT_BENT',
                'number': altBentMatch.group(1),
                'format': altBentMatch.group(2) or '',
                'content': displayText,
                'fullContent': altBentContent,  # ✅ Tam içerik
                'level': 4
            }
        
        return None


class ContentExtractor:
    """✅ TAM DOĞRU: ragChunkingTest.py'den"""
    
    @staticmethod
    def extract_between_headers(allLines: List[Dict], startIdx: int, currentLevel: int, currentHeader: Dict) -> Optional[str]:
        contentSegments = []
        
        firstLine = allLines[startIdx]
        
        if all(seg['isStrikethrough'] for seg in firstLine['segments']):
            return None
        
        if currentHeader['type'] == 'MADDE':
            maddeFullTitle = currentHeader.get('fullTitle', currentHeader['title'])
            maddeHeader = f"{currentHeader['number']} - {maddeFullTitle}"
            contentSegments.append({
                'text': maddeHeader,
                'isRed': False,
                'isStrikethrough': False
            })
        else:
            contentSegments.extend(firstLine['segments'])
        
        for i in range(startIdx + 1, len(allLines)):
            line = allLines[i]
            lineText = line['text'].strip()
            
            if not lineText:
                continue
            
            headerInfo = HeaderDetector.detect(lineText, line['segments'])
            
            if headerInfo and headerInfo['level'] <= currentLevel:
                break
            
            contentSegments.extend(line['segments'])
        
        mergedSegments = ParagraphSegmentExtractor.merge_consecutive(contentSegments)
        filteredSegments = [seg for seg in mergedSegments if not seg['isStrikethrough']]
        
        if not filteredSegments:
            return None
        
        plainContent = ''.join([seg['text'] for seg in filteredSegments])
        
        return plainContent.strip() if plainContent.strip() else None


class HierarchyTracker:
    def __init__(self):
        self.stack = []
        self.currentBolum = None
        self.currentMadde = None
    
    def setBolum(self, bolumText: str):
        self.currentBolum = bolumText
        self.currentMadde = None
        self.stack = []
    
    def setMadde(self, maddeNumber: str, maddeTitle: str):
        self.currentMadde = f"[MADDE {maddeNumber}] {maddeTitle}"
        self.stack = []
    
    def pushHeader(self, headerInfo: Dict, formattedTitle: str):
        level = headerInfo['level']
        self.stack = [item for item in self.stack if item['level'] < level]
        self.stack.append({
            'level': level,
            'title': formattedTitle,
            'type': headerInfo['type']
        })
    
    def getParent(self) -> Optional[str]:
        if len(self.stack) == 0:
            return None
        return self.stack[0]['title']
    
    def getCurrentMadde(self) -> Optional[str]:
        return self.currentMadde


class ChunkCreator:
    @staticmethod
    def format_content_with_commas(content: str) -> str:
        patterns = [
            r'\)\s*\d+\)',
            r'\)\s*[a-zçğıöşü]{1,3}\)',
            r'\)\s*\d+-',
            r'\)\s*[a-zçğıöşü]-',
        ]
        
        for pattern in patterns:
            content = re.sub(pattern, lambda m: m.group(0)[0] + ', ' + m.group(0)[1:].lstrip(), content)
        
        return content
    
    @staticmethod
    def create(headerInfo: Dict, plainContent: str, hierarchyTracker: HierarchyTracker) -> Optional[str]:
        if plainContent is None:
            return None
        
        plainContent = TextCleaner.clean_variants(plainContent)
        plainContent = TextCleaner.clean_price_formatting(plainContent)
        
        if not plainContent.strip():
            return None
        
        currentMadde = hierarchyTracker.getCurrentMadde()
        if currentMadde:
            maddeMatch = re.match(r'\[MADDE ([^\]]+)\]\s*(.+)', currentMadde)
            if maddeMatch:
                maddeNumber = maddeMatch.group(1)
                maddeTitle = maddeMatch.group(2)
                maddePrefix = f"{maddeNumber} - {maddeTitle}:"
                
                if not plainContent.startswith(maddeNumber):
                    plainContent = f"{maddePrefix} {plainContent}"
        
        parent = hierarchyTracker.getParent()
        if parent and parent not in plainContent:
            plainContent = f"{parent} > {plainContent}"
        
        plainContent = ChunkCreator.format_content_with_commas(plainContent)
        
        sutEkler = MetadataExtractor.extract_sut_ekler(plainContent)
        lawNumbers = MetadataExtractor.extract_law_numbers(plainContent)
        
        metadataParts = []
        if sutEkler:
            metadataParts.append(f"(SUT Ekleri: {', '.join(sutEkler)})")
        if lawNumbers:
            metadataParts.append(f"(Kanunlar: {', '.join(lawNumbers)})")
        
        if metadataParts:
            plainContent = f"{plainContent} {'. '.join(metadataParts)}."
        
        return plainContent


class SutDocx:
    def __init__(self):
        pass

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        
        allLines = []
        headers = []
        contentStarted = False
        listCounters = {}
        
        for paraIdx, para in enumerate(self.doc.paragraphs):
            segments = ParagraphSegmentExtractor.extract(para, listCounters, self.doc)
            
            if not segments:
                continue
            
            lineText = ''.join([seg['text'] for seg in segments])
            
            if not contentStarted:
                if RegexPatterns.BOLUM.match(lineText.strip()):
                    contentStarted = True
                else:
                    continue
            
            lineData = {
                'text': lineText,
                'segments': segments
            }
            
            allLines.append(lineData)
            
            headerInfo = HeaderDetector.detect(lineText, segments)
            
            if headerInfo:
                headerInfo['lineIndex'] = len(allLines) - 1
                headers.append(headerInfo)
        
        chunks = []
        hierarchyTracker = HierarchyTracker()
        
        for header in headers:
            headerType = header['type']
            
            if headerType == 'BOLUM':
                hierarchyTracker.setBolum(header['text'])
                continue
            
            elif headerType == 'MADDE':
                maddeTitleFull = header.get('fullTitle', header['title'])
                hierarchyTracker.setMadde(header['number'], maddeTitleFull)
                listCounters.clear()
                continue
            
            elif headerType == 'FIKRA':
                fikraContentFull = header.get('fullContent', header['content'])
                formattedTitle = f"({header['number']}) {fikraContentFull}"
                hierarchyTracker.pushHeader(header, formattedTitle)
            
            elif headerType == 'BENT':
                formatChar = header.get('format', ')')
                bentContentFull = header.get('fullContent', header['content'])
                formattedTitle = f"{header['number']}{formatChar} {bentContentFull}"
                hierarchyTracker.pushHeader(header, formattedTitle)
            
            elif headerType == 'ALT_BENT':
                formatChar = header.get('format', '')
                altBentContentFull = header.get('fullContent', header['content'])
                formattedTitle = f"{header['number']}{formatChar} {altBentContentFull}"
                hierarchyTracker.pushHeader(header, formattedTitle)
            
            else:
                continue
            
            plainContent = ContentExtractor.extract_between_headers(
                allLines, header['lineIndex'], header['level'], header
            )
            
            if plainContent is None:
                continue
            
            chunk = ChunkCreator.create(header, plainContent, hierarchyTracker)
            
            if chunk is None:
                continue
            
            chunks.append(chunk)
        
        return chunks


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Turkish", callback=None, **kwargs):
    """
    ✅ TAM DÜZELTİLMİŞ: ragChunkingTest.py mantığı
    """
    
    parser_config = kwargs.get("parser_config", {})
    
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    
    eng = lang.lower() == "english"
    
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "SUT dokumani isleniyor...")
        rawChunks = SutDocx()(filename, binary, from_page, to_page)
        callback(0.7, "SUT chunking tamamlandi.")
        
        result = []
        for chunkContent in rawChunks:
            d = copy.deepcopy(doc)
            d["content_with_weight"] = chunkContent
            d["content_ltks"] = rag_tokenizer.tokenize(chunkContent)
            d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
            result.append(d)
        
        callback(0.9, f"{len(result)} chunk olusturuldu.")
        return result
    
    else:
        raise NotImplementedError("SUT chunker sadece .docx dosyalarini destekler")


if __name__ == "__main__":
    import sys
    
    def dummy(prog=None, msg=""):
        pass
    
    chunk(sys.argv[1], callback=dummy)